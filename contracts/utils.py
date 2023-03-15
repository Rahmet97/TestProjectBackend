import os
import qrcode
import subprocess

from io import BytesIO
from django.http import HttpResponse
from django.template.loader import get_template, render_to_string

from xhtml2pdf import pisa

import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

from django.conf import settings
from rest_framework import validators, status

# from xhtml2pdf.default import DEFAULT_CSS


def error_response_404():
    raise validators.ValidationError(
        detail={"message": "Object is not found 404"}, code=status.HTTP_404_NOT_FOUND)

def error_response_500():
    raise validators.ValidationError(
        detail={"message": "Internal server error 500"}, code=500)


# Numbers to word
class NumbersToWord:

    # Digit Numbers to word: 21 -> yigirma bir
    def _hundreds(self, m):
        digits = {
            1: "bir", 2: "ikki", 3: "uch", 4: "to'rt", 5: "besh", 6: "olti", 7: "yetti", 8: "sakkiz", 9: "to'qqiz",
            10: "o'n", 20: "yigirma", 30: "o'ttiz", 40: "qirq", 50: "ellik", 60: "oltmish", 70: "yetmish", 80: "sakson",
            90: "to'qson"
        }

        d1 = m // 100
        d2 = (m // 10) % 10
        d3 = m % 10
        s1, s2, s3 = '', '', ''
        if d1 != 0:
            s1 = f'{digits[d1]} yuz '
        if d2 != 0:
            s2 = digits[d2 * 10] + ' '
        if d3 != 0:
            s3 = digits[d3] + ' '

        return s1 + s2 + s3

    # Numbers to word
    def _number2word(self, n):
        d1 = {0: "", 1: "ming ", 2: "million ", 3: "milliard ", 4: "trillion "}

        fraction = []
        while n > 0:
            r = n % 1000
            fraction.append(r)
            n //= 1000

        s = ''
        for i in range(len(fraction)):
            if fraction[i] != 0:
                yuz = self._hundreds(fraction[i]) + d1[i]
                s = yuz + s

        return s.rstrip()
    
    # change number to word
    def change_num_to_word(self, n: int) -> str:
        return self._number2word(n=n) 


# create qr code
def create_qr(link):
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_H,
        box_size=10,
        border=4
    )
    file_name = link.split('/')[-1]
    file_path = f"{settings.MEDIA_ROOT}/qr/"
    qr.add_data(link)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white").convert('RGB')
    # img.save(file_path + file_name.split('.')[0] + '.png')
    img.save(file_path + file_name.split('/')[-1] + '.png')
    
    print("qr_file_name: ", file_name.split('/')[-1] + '.png')
    print("qr_file_path1: ", file_path + file_name.split('/')[-1] + '.png')
    print("qr_file_path2: ", file_path + file_name.split('.')[0] + '.png')


    return file_path + file_name.split('.')[0] + '.png'


def set_docx_font_style(docx_file_path):
    # Open the DOCX file with python-docx
    doc = docx.Document(docx_file_path)

    # Create a new style based on the existing Normal style with Times New Roman font
    times_new_roman_style = doc.styles.add_style('Times New Roman', WD_STYLE_TYPE.PARAGRAPH)
    font = times_new_roman_style.font
    font.name = 'Times New Roman'

    # Apply the Times New Roman style to all paragraphs in the document
    for paragraph in doc.paragraphs:
        paragraph.style = times_new_roman_style

    # Save the modified document to a new file
    new_docx_file_path = os.path.splitext(docx_file_path)[0] + '_modified.docx'
    doc.save(new_docx_file_path)

    return new_docx_file_path


def convert_docx_to_pdf(docx_file_path: str):
    """
    Convert a DOCX file to PDF using LibreOffice.

    :param docx_file_path: str, the path to the DOCX file
    :return: str, the path to the converted PDF file
    """
    docx_file_path = set_docx_font_style(docx_file_path)

    path = "/".join(docx_file_path.split('/')[0:-1]) + '/'
    pdf_file_path = f"{path}{docx_file_path.split('/')[-1].split('.')[0]}.pdf"

    # Create the command to convert DOCX to PDF using libreoffice
    # command = ['libreoffice', '--headless', '--convert-to', 'pdf', docx_file_path, '--outdir', path]
    # command = ['libreoffice', '--headless', '--convert-to', 'pdf:writer_pdf_Export', docx_file_path, '--outdir', path, '--nofirststartwizard', '--nolockcheck', '--nologo', '--nodisplay', '--norestore', '--convert-images-to-jpg', '--writer-pdf-embedfonts', '--writer-pdf-use-cff', '--writer-pdf-subset-fonts']    
    # command = ['libreoffice', '--headless', '--convert-to', 'pdf', '--convert-to', 'pdf:writer_pdf_Export', '--outdir', path, '--nofirststartwizard', '--nolockcheck', '--nologo', '--nodisplay', '--norestore', '--convert-images-to-jpg', '--writer-pdf-embedfonts', '--writer-pdf-use-cff', '--writer-pdf-subset-fonts', docx_file_path]    
    command = ['libreoffice', '--headless', '--convert-to', 'pdf:writer_pdf_Export', docx_file_path, '--outdir', path, '--nofirststartwizard', '--nolockcheck', '--nologo', '--norestore']    


    # Run the command in the terminal using subprocess with utf-8 encoding and capture the output
    result = subprocess.run(command, encoding='utf-8', stdout=subprocess.PIPE, stderr=subprocess.PIPE)

    # Print the output of the command
    print("result.stdout: ", result.stdout)

    # subprocess.call(command)
    print("pdf", pdf_file_path)
    return pdf_file_path


def render_to_pdf(template_src: str, context_dict={}):
    template = get_template(template_name=template_src)
    html = template.render(context_dict)
    
    result = BytesIO()
    # pdf = pisa.pisaDocument(BytesIO(html.encode("ISO-8859-1")), result)
    pdf = pisa.pisaDocument(BytesIO(html.encode("utf-8")), result)

    if not pdf.err:
        pdf_content_with_page_breaks = insert_page_breaks(result.getvalue())
        return HttpResponse(pdf_content_with_page_breaks, content_type='application/pdf')
    return None

def insert_page_breaks(pdf_content):
    page_break_tag = '<div style="page-break-after: always;"></div>'
    pdf_content_with_page_breaks = pdf_content.replace(b'<body>', b'<body>' + page_break_tag.encode('utf-8'))
    return pdf_content_with_page_breaks


# def render_to_pdf(template_src: str, context_dict={}):
#     html = render_to_string(template_name=template_src, context=context_dict)

#     # Read the CSS file into a string
#     with open("/usr/src/app/static/shablon/shablon.css", "r") as css_file:
#         css = css_file.read()
    
#     # Generate a PDF file from the HTML and CSS using xhtml2pdf
#     response = HttpResponse(content_type="application/pdf")
#     response["Content-Disposition"] = 'attachment; filename="my_pdf.pdf"'
#     pdf = pisa.CreatePDF(html, dest=response, encoding="utf-8", css=css)
#     if pdf.err:
#         return HttpResponse("Error generating PDF file.")
#     return response


# def render_to_pdf(template_src: str, context_dict={}):
#     template = get_template(template_name=template_src)
#     html = template.render(context_dict)

#     result = BytesIO()
#     pdf = pisa.CreatePDF(
#         BytesIO(html.encode("utf-8")),
#         result,
#         encoding='utf-8',
#         show_error_as_pdf=True,
#         link_callback=None,
#         debug=0,
#         path='',
#         default_css=None,
#         user_css=None,
#         font_config=None,
#         all_texts=False,
#         xhtml=False,
#         xml_output=None,
#         keep_tmp=False,
#         strict=False,
#         raise_exception=True,
#         pisa_context=pisa.PisaContext()
#     )

#     if not pdf.err:
#         pdf_content_with_page_breaks = insert_page_breaks(result.getvalue())
#         return HttpResponse(pdf_content_with_page_breaks, content_type='application/pdf')
#     return None


# def insert_page_breaks(pdf_content):
#     page_break_tag = '<hr style="page-break-before: always;">'
#     pdf_content_with_page_breaks = pdf_content.replace(b'<body>', b'<body>' + page_break_tag.encode('utf-8'))
#     return pdf_content_with_page_breaks


def delete_file(file_path: str):
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
            print("The file has been deleted.")
        except Exception as e:
            print(e)
    else:
        print("The file does not exist.")
